import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.metrics import classification_report, confusion_matrix, ConfusionMatrixDisplay
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from docx import Document
from docx.shared import Inches
import joblib

# Carrega e lê o banco de dados
df = pd.read_csv("churn.csv")

# Mostra a estrutura do arquivo
print(df.shape)
print(df.columns.values)
print(df.isna().sum())
print(df.describe())

# Calcula e mostra a % de Churn
numRetained = df[df.Churn == 'No'].shape[0]
numChurned = df[df.Churn == 'Yes'].shape[0]
percent_retained = numRetained / (numRetained + numChurned) * 100
percent_churned = numChurned / (numRetained + numChurned) * 100
print(f"{percent_retained:.2f}% of customers stayed in the company")
print(f"{percent_churned:.2f}% of customers left")

# Essa etapa cria um DOCx para exportar os resultados e facilitar a análise
doc = Document()

# Plot de Churn por gênero
sns.countplot(x='gender', hue='Churn', data=df)
plt.savefig("churn_by_gender.png")
plt.show()
doc.add_heading('Churn by Gender', level=1)
doc.add_picture('churn_by_gender.png', width=Inches(6))

# Plot de Churn causado pelo serviço 'internet'
sns.countplot(x='InternetService', hue='Churn', data=df)
plt.savefig("churn_by_internet_service.png")
plt.show()
doc.add_heading('Churn by Internet Service', level=1)
doc.add_picture('churn_by_internet_service.png', width=Inches(6))

# Plot numérico de Churn por mês
numericFeatures = ['tenure', 'MonthlyCharges']
fig, ax = plt.subplots(1, 2, figsize=(28, 8))
df[df.Churn == "No"][numericFeatures].hist(bins=20, color='blue', alpha=0.5, ax=ax)
df[df.Churn == "Yes"][numericFeatures].hist(bins=20, color='orange', alpha=0.5, ax=ax)
plt.savefig("churn_by_numeric_features.png")
plt.show()
doc.add_heading('Churn by Numeric Features', level=1)
doc.add_picture('churn_by_numeric_features.png', width=Inches(6))

# Data cleaning e Pré-processamento
cleanDF = df.drop('customerID', axis=1)
for column in cleanDF.columns:
    if cleanDF[column].dtype == 'object':
        cleanDF[column] = LabelEncoder().fit_transform(cleanDF[column])

# Escalonamento da feature
x = cleanDF.drop('Churn', axis=1)
y = cleanDF['Churn']
x = StandardScaler().fit_transform(x)

# Dividindo os dados
xtrain, xtest, ytrain, ytest = train_test_split(x, y, test_size=0.2, random_state=42)

# Modelo de treinamento da IA de 'Regressão Logistica'
model = LogisticRegression(max_iter=1000)
model.fit(xtrain, ytrain)

# Salva o modelo treinado
joblib.dump(model, 'churn_model.pkl')

# Precisão do modelo de IA
predictions = model.predict(xtest)

# Mostra a precisão do modelo
print(predictions)

# Mostra a classificação dos valores do modelo
report = classification_report(ytest, predictions)
print(report)

# Adiciona no documento criado, a classificação e os plots
doc.add_heading('Classification Report', level=1)
doc.add_paragraph(report)

# Gera a matriz de confusão
conf_matrix = confusion_matrix(ytest, predictions)

# Exibe a matriz de confusão
disp = ConfusionMatrixDisplay(confusion_matrix=conf_matrix, display_labels=model.classes_)
disp.plot()
plt.savefig("confusion_matrix.png")
plt.show()
doc.add_heading('Confusion Matrix', level=1)
doc.add_picture('confusion_matrix.png', width=Inches(6))

# Salva a edição final do doc
doc.save("churn_analysis.docx")

# Criar um arquivo de exemplo new_customers.csv
new_customers = df.drop(['Churn'], axis=1).sample(10, random_state=42)
new_customers.to_csv('new_customers.csv', index=False)

# Carregar o modelo salvo e fazer previsões em novos dados
def predict_new_data(new_data_path):
    # Carregue o modelo
    model = joblib.load('churn_model.pkl')

    # Carregue os novos dados
    new_data = pd.read_csv(new_data_path)
    new_data_clean = new_data.drop('customerID', axis=1)

    for column in new_data_clean.columns:
        if new_data_clean[column].dtype == 'object':
            new_data_clean[column] = LabelEncoder().fit_transform(new_data_clean[column])

    new_data_scaled = StandardScaler().fit_transform(new_data_clean)

    # Faça previsões
    new_predictions = model.predict(new_data_scaled)

    # Exiba as previsões
    return new_predictions, new_data

# Exemplo de uso da função de previsão
# Previsões em novos dados
new_predictions, new_data = predict_new_data('new_customers.csv')
print(new_predictions)

# Comparação e exportação para docx
def export_predictions_to_doc(new_data, new_predictions):
    # Cria um novo documento ou abre um existente
    doc = Document('churn_analysis.docx')
    doc.add_heading('New Customer Predictions', level=1)

    for i in range(len(new_data)):
        doc.add_paragraph(f"Customer {new_data.iloc[i]['customerID']}: Predicted Churn - {'Yes' if new_predictions[i] == 1 else 'No'}")

    # Calcula a porcentagem de churn prevista
    predicted_churn_rate = np.mean(new_predictions) * 100
    predicted_retained_rate = 100 - predicted_churn_rate

    # Adiciona a comparação das previsões com os valores reais
    doc.add_heading('Comparison of Predictions with Exploratory Analysis', level=1)
    doc.add_paragraph(f"Percentage of customers predicted to stay: {predicted_retained_rate:.2f}%")
    doc.add_paragraph(f"Percentage of customers predicted to churn: {predicted_churn_rate:.2f}%")
    doc.add_paragraph(f"Percentage of customers actually stayed (exploratory): {percent_retained:.2f}%")
    doc.add_paragraph(f"Percentage of customers actually churned (exploratory): {percent_churned:.2f}%")

    # Salva o documento
    doc.save('churn_analysis.docx')

# Exportar as previsões para o documento
export_predictions_to_doc(new_data, new_predictions)

